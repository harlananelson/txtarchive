# txtarchive/word_converter.py
"""
Word document to Markdown converter for txtarchive.
Supports .docx files and converts them to clean markdown format.
"""

import os
from pathlib import Path
from .header import logger
import zipfile
import xml.etree.ElementTree as ET
import re

# Optional dependencies - will be imported if available
try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False
    logger.warning("mammoth not available - basic conversion only")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not available - limited Word support")

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

def convert_docx_with_mammoth(docx_path):
    """
    Convert Word document to markdown using mammoth library.
    This provides the best formatting preservation.
    """
    if not MAMMOTH_AVAILABLE:
        raise ImportError("mammoth library is required for high-quality conversion")
    
    try:
        with open(docx_path, "rb") as docx_file:
            # Convert to HTML first, then to markdown
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
            
            # Convert HTML to markdown (basic conversion)
            markdown_content = html_to_markdown(html_content)
            
            # Log any warnings
            if result.messages:
                for message in result.messages:
                    logger.warning(f"Mammoth warning: {message}")
            
            return markdown_content
            
    except Exception as e:
        logger.error(f"Error converting with mammoth: {e}")
        raise

def convert_docx_with_python_docx(docx_path):
    """
    Convert Word document to markdown using python-docx library.
    Provides basic text extraction with paragraph structure.
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx library is required")
    
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_content.append("")
                continue
            
            # Detect headings based on style
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name:
                # Extract heading level
                level = 1
                if 'heading 1' in style_name:
                    level = 1
                elif 'heading 2' in style_name:
                    level = 2
                elif 'heading 3' in style_name:
                    level = 3
                elif 'heading 4' in style_name:
                    level = 4
                elif 'heading 5' in style_name:
                    level = 5
                elif 'heading 6' in style_name:
                    level = 6
                
                markdown_content.append(f"{'#' * level} {text}")
            elif 'title' in style_name:
                markdown_content.append(f"# {text}")
            else:
                markdown_content.append(text)
        
        # Process tables
        for table in doc.tables:
            markdown_content.append("")  # Empty line before table
            
            # Header row
            if table.rows:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                markdown_content.append("| " + " | ".join(header_cells) + " |")
                markdown_content.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                
                # Data rows
                for row in table.rows[1:]:
                    data_cells = [cell.text.strip() for cell in row.cells]
                    markdown_content.append("| " + " | ".join(data_cells) + " |")
            
            markdown_content.append("")  # Empty line after table
        
        return "\n".join(markdown_content)
        
    except Exception as e:
        logger.error(f"Error converting with python-docx: {e}")
        raise

def convert_docx_with_pandoc(docx_path):
    """
    Convert Word document to markdown using pandoc.
    Requires pandoc to be installed on the system.
    """
    if not PANDOC_AVAILABLE:
        raise ImportError("pypandoc library is required")
    
    try:
        # Convert using pandoc
        markdown_content = pypandoc.convert_file(
            docx_path, 
            'markdown', 
            format='docx',
            extra_args=['--wrap=none', '--markdown-headings=atx']
        )
        return markdown_content
        
    except Exception as e:
        logger.error(f"Error converting with pandoc: {e}")
        raise

def basic_docx_extraction(docx_path):
    """
    Basic Word document text extraction by parsing the XML directly.
    Fallback method when no specialized libraries are available.
    """
    try:
        markdown_content = []
        
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # Extract main document XML
            document_xml = docx_zip.read('word/document.xml')
            
            # Parse XML
            root = ET.fromstring(document_xml)
            
            # Namespace mapping for Word XML
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # Extract paragraphs
            for para in root.findall('.//w:p', namespaces):
                text_elements = para.findall('.//w:t', namespaces)
                paragraph_text = ''.join(elem.text or '' for elem in text_elements)
                
                if paragraph_text.strip():
                    markdown_content.append(paragraph_text.strip())
                else:
                    markdown_content.append("")  # Empty line
        
        return "\n".join(markdown_content)
        
    except Exception as e:
        logger.error(f"Error with basic extraction: {e}")
        raise

def html_to_markdown(html_content):
    """Convert HTML to markdown, delegating to the shared html_converter module."""
    if not html_content:
        return ""
    from .html_converter import convert_html_to_markdown_text
    return convert_html_to_markdown_text(html_content)

def convert_word_to_markdown(docx_path, method='auto'):
    """
    Convert a Word document to markdown format.
    
    Args:
        docx_path (str or Path): Path to the .docx file
        method (str): Conversion method - 'auto', 'mammoth', 'docx', 'pandoc', or 'basic'
    
    Returns:
        str: Markdown content
    """
    docx_path = Path(docx_path)
    
    if not docx_path.exists():
        raise FileNotFoundError(f"Word document not found: {docx_path}")
    
    if not docx_path.suffix.lower() == '.docx':
        raise ValueError(f"Only .docx files are supported, got: {docx_path.suffix}")
    
    logger.info(f"Converting Word document to markdown: {docx_path}")
    
    if method == 'auto':
        # Try methods in order of preference
        if MAMMOTH_AVAILABLE:
            method = 'mammoth'
        elif PANDOC_AVAILABLE:
            method = 'pandoc'
        elif PYTHON_DOCX_AVAILABLE:
            method = 'docx'
        else:
            method = 'basic'
        
        logger.info(f"Auto-selected conversion method: {method}")
    
    try:
        if method == 'mammoth':
            return convert_docx_with_mammoth(docx_path)
        elif method == 'pandoc':
            return convert_docx_with_pandoc(docx_path)
        elif method == 'docx':
            return convert_docx_with_python_docx(docx_path)
        elif method == 'basic':
            return basic_docx_extraction(docx_path)
        else:
            raise ValueError(f"Unknown conversion method: {method}")
            
    except ImportError as e:
        logger.error(f"Required library not available for method '{method}': {e}")
        if method != 'basic':
            logger.info("Falling back to basic extraction")
            return basic_docx_extraction(docx_path)
        else:
            raise

def convert_word_documents_in_directory(directory, output_dir=None, method='auto', replace_existing=False):
    """
    Convert all Word documents in a directory to markdown files.
    
    Args:
        directory (str or Path): Directory containing .docx files
        output_dir (str or Path): Output directory for markdown files (default: same as input)
        method (str): Conversion method
        replace_existing (bool): Whether to overwrite existing markdown files
    
    Returns:
        list: List of (input_path, output_path) tuples for successful conversions
    """
    directory = Path(directory)
    if output_dir is None:
        output_dir = directory
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    
    converted_files = []
    
    for docx_file in directory.rglob("*.docx"):
        if docx_file.name.startswith("~$"):  # Skip temporary files
            continue
            
        # Create output path
        relative_path = docx_file.relative_to(directory)
        output_path = output_dir / relative_path.with_suffix('.md')
        
        # Create output directory if needed
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Check if output already exists
        if output_path.exists() and not replace_existing:
            logger.info(f"Skipping existing file: {output_path}")
            continue
        
        try:
            logger.info(f"Converting: {docx_file} -> {output_path}")
            markdown_content = convert_word_to_markdown(docx_file, method=method)
            
            # Write markdown file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            converted_files.append((docx_file, output_path))
            logger.info(f"Successfully converted: {output_path}")
            
        except Exception as e:
            logger.error(f"Failed to convert {docx_file}: {e}")
            continue
    
    logger.info(f"Converted {len(converted_files)} Word documents to markdown")
    return converted_files

